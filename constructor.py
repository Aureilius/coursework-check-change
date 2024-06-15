from docx import Document
from docx.shared import Pt
from docx.dml.color import ColorFormat



def construct():
    doc = Document()

    try:
        part_cnt = int(input('Введите количество частей курсовой: '))
    except:
        part_cnt = 0
    if part_cnt > 3:
        raise 'Слишком много частей!'
    part_names = []
    subsections = []
    subsections_names = []
    for i in range(part_cnt):
        part_names.append(input(f'Введите название {i+1} части: '))
        try:
            subsections.append(int(input(f'Введите количество подразделов курсовой {i+1} части: ')))
        except:
            subsections.append(0)
        if subsections[-1] > 10:
            raise 'Слишком много подпунктов!'
        for j in range(subsections[i]):
            subsections_names.append(input(f'Введите название подраздела курсовой {i+1}.{j+1}: '))

    #Содержание
    doc.add_paragraph('Содержание\n').paragraph_format.alignment = 1
    doc.add_paragraph('Введение').paragraph_format.alignment = 2
    for i in range(part_cnt):
        doc.add_paragraph(f'{i + 1} {part_names[i]}\n'.upper()).paragraph_format.alignment = 2
        for x in range(subsections[i]):
            doc.add_paragraph(f'{i + 1}.{x + 1} {subsections_names[x]}\n').paragraph_format.alignment = 2
    doc.add_paragraph('Заключение\nСписок библиографических источников').paragraph_format.alignment = 2
    doc.add_page_break()

    #Введение
    doc.add_paragraph('Введение\n').paragraph_format.alignment = 1
    print('Введите текст введения: Когда всё напишите, напишите stop с новой строки для продолжения')
    while True:
        line = input()
        if line == 'stop':
            break
        doc.add_paragraph(f'\t{line}')
    doc.add_page_break()


    #Написание частей
    for i in range(part_cnt):
        doc.add_paragraph(f'{i+1} {part_names[i]}\n'.upper()).paragraph_format.alignment = 1
        for x in range(subsections[i]):
            doc.add_paragraph(f'{i+1}.{x+1} {subsections_names[x]}\n').paragraph_format.alignment = 1
            print(f'Введите текст подраздела {i + 1}.{x + 1} {subsections_names[x]}:\nКогда всё напишите, напишите stop с новой строки для продолжения')
            while True:
                line = input()
                if line == 'stop':
                    break
                doc.add_paragraph(f'\t{line}')
            doc.add_page_break()



    #Заключение
    doc.add_paragraph('Заключение\n').paragraph_format.alignment = 1
    print('Введите текст заключения: Когда всё напишите, напишите stop с новой строки для продолжения')
    while True:
        line = input()
        if line == 'stop':
            break
        doc.add_paragraph(f'\t{line}')
    doc.add_page_break()

    #Список источников
    doc.add_paragraph('Список библиографических источников\n').paragraph_format.alignment = 1
    print('Введите список библиографических источников.\nНовый источник с новой строки\nКогда всё напишите, напишите stop с новой строки для продолжения')
    content = []
    while True:
        line = input()
        if line == 'stop':
            break
        content.append(line)
    content.sort()
    for i in range(len(content)):
        s = f'{i+1}. {content[i].capitalize()}'
        doc.add_paragraph(s)

    #Работа над оформлением
    for p in doc.paragraphs:
        p.paragraph_format.space_before = p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing = 1.5
        if p.alignment == 0:
            p.alignment = 3
            p.paragraph_format.first_line_indent = Pt(18)
            for run in p.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(14)
                font.underline = None
                font.highlight_color = None
                if font.bold and font.italic:
                    font.italic = True
                    font.bold = False
        if p.alignment == 1:
            for run in p.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(14)
                font.underline = None
                font.highlight_color = None
                font.bold = True
    for p in doc.paragraphs:
        if p.alignment == 2:
            p.alignment = 0
        for run in p.runs:
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(14)
            font.underline = None
            font.highlight_color = None


    #Сохранение документа
    doc.save('result.docx')
    print('Документ сохранен в директории программы. Удачи!')