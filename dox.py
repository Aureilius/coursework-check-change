from docx import Document
from docx.shared import Pt

def pointer(path):
    res = int(input('Что требуется сделать с файлом:\n1. Проверить на оформление\n2. Подогнать под оформление\nВаш ответ: '))
    if res == 1:
        check(path)
    elif res == 2:
        change(path)

def change(path):
    doc = Document(path)
    for p in doc.paragraphs:
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing = 1.5
        p.alignment = 3
        for run in p.runs:
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(14)
            font.underline = None
            font.highlight_color = None
    for s in doc.sections:
        s.orientation = 0
    doc.save('result.docx')
    print('Документ сохранен в директории программы. Удачи!')

def check(path):
    doc = Document(path)

    good = "Соответствует"
    bad = "Не соответствует"

    font_check = good
    font_size_check = good
    font_underline = good
    font_color = good
    line_space = good
    aligment_check = good
    space_paragraph = good
    orientation = good
    highlight = good
    first_intend = good
    bold_italic = good

    for p in doc.paragraphs:
        if p.paragraph_format.space_before != 0 or p.paragraph_format.space_after != 0:
            space_paragraph = bad
        if p.paragraph_format.line_spacing != 1.5:
            line_space = bad
        if p.paragraph_format.first_line_indent and p.alignment != 3:
            first_intend = bad
        for run in p.runs:
            font = run.font
            if p.alignment == 1 and not(font.bold):
                aligment_check = bad
            if font.name != 'Times New Roman':
                font_check = bad
            if font.size != Pt(14):
                font_size_check = bad
            if font.underline:
                font_underline = 'Уберите подчеркивание текста!'
            if font.color.rgb:
                font_color = bad
            if font.highlight_color:
                highlight = bad
            if font.bold and font.italic:
                bold_italic = 'Нельзя одновременно использовать жирный шрифт и курсив!'
    for s in doc.sections:
        if s.orientation != 0:
            orientation = bad

    print(f"Стиль шрифта: {font_check}\n"
          f"Размер шрифта: {font_size_check}\n"
          f"Отсутствие подчеркивания шрифта: {font_underline}\n"
          f"Цвет шрифта: {font_color}\n"
          f"Межстрочный интервал: {line_space}\n"
          f"Расположение текста: {aligment_check}\n"
          f"Пробел до/после абзацев: {space_paragraph}\n"
          f"Ориентация страницы: {orientation}\n"
          f"Цвет заднего фона текста: {highlight}\n"
          f"Отступ в начале абзаца: {first_intend}\n"
          f"Отсутствие жирного шрифта и курсива одновременно: {bold_italic}\n")